"""
Generates the academic article as a formatted Word document.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FIGURES = os.path.join('reports', 'figures')
OUT = 'Article_Academique_Detection_FakeNews_COVID19.docx'

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Helper functions ──────────────────────────────────────────────────────────
def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x00, 0x4A, 0x80)
    return p

def para(text, bold=False, italic=False, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    return p

def add_figure(path, caption, width=5.5):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.runs[0]
    run.font.size = Pt(9)
    run.font.italic = True

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(40)
run = title_p.add_run(
    "Détection Automatique de Fausses Informations liées au COVID-19 :\n"
    "Comparaison de Modèles d'Apprentissage Automatique Classiques\n"
    "et Transformers avec Analyse d'Explicabilité"
)
run.font.size = Pt(16)
run.font.bold = True
run.font.color.rgb = RGBColor(0x00, 0x3A, 0x6E)

doc.add_paragraph()
auth = doc.add_paragraph()
auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
auth.add_run("Youssef Benmabrouk\n").font.size = Pt(12)
auth.add_run("École Polytechnique de Tunisie — Niveau 4CS\n").font.size = Pt(11)
auth.add_run("Encadrant : Mme Nadia\n").font.size = Pt(11)
auth.add_run("Mai 2026").font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# RÉSUMÉ
# ═══════════════════════════════════════════════════════════════════════════════
heading("Résumé", level=1)

para(
    "La pandémie de COVID-19 a donné naissance à un phénomène parallèle, "
    "qualifié d'« infodémie » par l'OMS : une propagation massive et incontrôlée "
    "de fausses informations sur les réseaux sociaux. Face à cette menace, nous "
    "avons développé dans le cadre de ce projet un système de classification "
    "automatique capable de distinguer les contenus véridiques des contenus "
    "trompeurs. Notre approche repose sur la comparaison rigoureuse de cinq "
    "modèles d'apprentissage automatique, allant des méthodes classiques (SVM, "
    "Régression Logistique, Random Forest, XGBoost) jusqu'aux architectures "
    "Transformer de dernière génération (DistilBERT). Nous avons travaillé sur "
    "le jeu de données Constraint COVID-19, constitué de 10 700 publications "
    "issues des réseaux sociaux. Le meilleur modèle obtenu est un SVM couplé à "
    "une représentation TF-IDF, atteignant 93,9 % de précision et un score F1 "
    "de 0,939. DistilBERT, fine-tuné sur le même corpus, atteint 92,9 % avec "
    "une AUC-ROC de 0,976. Au-delà des performances brutes, nous avons intégré "
    "des méthodes d'explicabilité — LIME et visualisation des poids d'attention — "
    "afin de rendre les décisions du modèle interprétables. Cette analyse révèle "
    "des biais linguistiques et culturels que tout système de modération de "
    "contenu devrait prendre en compte avant déploiement."
)

doc.add_paragraph()
kw_p = doc.add_paragraph()
kw_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
kw_p.add_run("Mots-clés : ").font.bold = True
kw_p.add_run(
    "Détection de fake news, COVID-19, Traitement du Langage Naturel, "
    "DistilBERT, SVM, TF-IDF, LIME, Explicabilité, Réseaux sociaux, Infodémie."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════
heading("1. Introduction", level=1)

para(
    "Quand le SARS-CoV-2 a commencé à se propager à travers le monde à la fin "
    "de 2019, il n'a pas fallu longtemps avant qu'une deuxième crise émerge, "
    "cette fois-ci informationnelle. L'OMS a officiellement parlé d'« infodémie » "
    "dès le mois de février 2020 [1], un terme qui désigne la surabondance "
    "d'informations — exactes ou fausses — qui rend difficile, voire impossible, "
    "de trouver des sources fiables au moment où on en a le plus besoin. Les "
    "plateformes de réseaux sociaux comme Twitter, WhatsApp et Facebook sont "
    "devenues des vecteurs de diffusion rapide pour toutes sortes d'affirmations "
    "non vérifiées : des théories du complot sur l'origine du virus, des "
    "« remèdes miracles » comme la Chloroquine, ou encore des rumeurs sur les "
    "effets des vaccins. Selon Sharma et al. [2], des millions de publications "
    "contenant des informations médicales erronées circulaient chaque jour en "
    "ligne pendant les premières semaines de la pandémie."
)

para(
    "Le fact-checking manuel, bien qu'essentiel, présente des limites évidentes "
    "en termes de volume et de réactivité. Une équipe d'analystes ne peut pas "
    "physiquement vérifier des dizaines de milliers de publications par heure. "
    "C'est dans ce contexte que les méthodes d'apprentissage automatique "
    "appliquées au Traitement du Langage Naturel (NLP) prennent tout leur sens : "
    "elles permettent d'automatiser une partie de ce travail de vérification, "
    "en identifiant les patterns linguistiques caractéristiques du contenu trompeur."
)

para(
    "Dans ce projet, nous avons abordé cette problématique sous deux angles "
    "complémentaires. D'abord, un angle purement technique : quelle architecture "
    "de modèle donne les meilleures performances sur un corpus COVID-19 spécifique ? "
    "Ensuite, un angle lié à la sécurité de l'IA : comment s'assurer que les "
    "décisions du modèle sont transparentes, auditables et exemptes de biais ? "
    "Cette deuxième question est particulièrement importante dans un contexte de "
    "santé publique, où une fausse alerte ou un silence sur un contenu dangereux "
    "peut avoir des conséquences réelles."
)

para(
    "Les contributions principales de ce travail sont au nombre de cinq. "
    "Premièrement, nous établissons des baselines solides avec quatre classifieurs "
    "classiques couplés à TF-IDF. Deuxièmement, nous fine-tunons DistilBERT et "
    "analysons ses courbes d'apprentissage. Troisièmement, nous appliquons LIME "
    "et la visualisation des poids d'attention pour rendre chaque prédiction "
    "interprétable. Quatrièmement, nous conduisons une analyse systématique des "
    "erreurs qui identifie la longueur du texte comme facteur principal d'échec. "
    "Cinquièmement, l'ensemble du pipeline est empaqueté dans une application "
    "web interactive permettant des inférences en temps réel."
)

para(
    "La suite de cet article est organisée comme suit. La section 2 présente les "
    "travaux connexes. La section 3 détaille notre méthodologie (dataset, "
    "prétraitement, architectures, évaluation). La section 4 présente et discute "
    "les résultats expérimentaux. La section 5 conclut et propose des perspectives."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 2. ÉTAT DE L'ART
# ═══════════════════════════════════════════════════════════════════════════════
heading("2. État de l'art", level=1)

para(
    "La détection automatique de fake news a émergé comme un champ de recherche "
    "actif au sein de la communauté NLP à partir de 2017, notamment avec la "
    "publication du FakeNewsChallenge dataset [8] qui a établi les premières "
    "références de performance. Les premières approches reposaient principalement "
    "sur des représentations bag-of-words couplées à des classifieurs classiques "
    "tels que SVM ou Naive Bayes [3]. Ces méthodes, bien que simples, obtenaient "
    "des résultats étonnamment compétitifs sur des benchmarks bien définis."
)

para(
    "Une deuxième vague d'approches a cherché à enrichir les représentations "
    "textuelles avec des features stylistiques et linguistiques : distributions "
    "de longueur de phrases, indices de lisibilité, ratios partie-du-discours, "
    "polarité émotionnelle [9]. Volkova et al. [10] ont montré que les marqueurs "
    "linguistiques d'incertitude épistémique (formules de hedging, verbes modaux) "
    "étaient systématiquement plus présents dans les contenus non vérifiés. "
    "Ces observations rejoignent ce que nous avons nous-mêmes constaté lors de "
    "notre exploration des données : les fake news COVID-19 tendent à employer "
    "un vocabulaire assertif et alarmiste, tandis que les nouvelles réelles "
    "utilisent davantage un registre institutionnel et mesuré."
)

para(
    "Le vrai tournant est venu avec BERT de Devlin et al. [4], publié en 2018. "
    "En pré-entraînant un transformeur profond et bidirectionnel sur d'énormes "
    "corpus non étiquetés, puis en l'affinant sur des tâches spécifiques, BERT "
    "a établi de nouveaux records sur de nombreux benchmarks NLP. Plusieurs "
    "équipes ont rapidement transposé cette architecture à la détection de fake "
    "news. Kula et al. [11] ont notamment rapporté des gains de 4 à 6 points "
    "de pourcentage par rapport aux meilleurs résultats antérieurs sur le LIAR "
    "dataset."
)

para(
    "La pandémie de COVID-19 a motivé une ligne de recherche spécifique. "
    "L'atelier Constraint à AAAI 2021 [13], qui a produit le jeu de données "
    "utilisé dans notre étude, a organisé la première tâche partagée de "
    "détection de fake news COVID-19. Les systèmes gagnants combinaient des "
    "représentations BERT avec des features lexicales issues de bases de "
    "connaissances externes, atteignant des scores F1 dans la fourchette "
    "97–98 %. Ces chiffres doivent toutefois être interprétés avec prudence, "
    "car ils reflètent des conditions d'évaluation sur des ensembles de test "
    "bien équilibrés qui ne reflètent pas nécessairement la diversité des "
    "contenus rencontrés en conditions réelles."
)

para(
    "Du côté de l'explicabilité, Popat et al. [14] ont appliqué LIME à des "
    "classifieurs basés sur les transformeurs et ont montré que les mots "
    "associés au sensationnalisme et aux théories du complot conduisaient "
    "systématiquement les prédictions vers la classe « fake ». La visualisation "
    "des poids d'attention dans les modèles de type BERT a été étudiée par "
    "Jain et Wallace [15], qui soulèvent des mises en garde importantes : "
    "un poids d'attention élevé n'est pas nécessairement synonyme d'importance "
    "causale. Malgré cela, les cartes d'attention restent un outil de diagnostic "
    "largement utilisé pour l'analyse qualitative, et c'est dans cet esprit "
    "que nous les utilisons dans ce travail."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 3. MÉTHODOLOGIE
# ═══════════════════════════════════════════════════════════════════════════════
heading("3. Méthodologie", level=1)

heading("3.1 Jeu de données", level=2)

para(
    "Nos expériences sont conduites sur le jeu de données Constraint COVID-19 "
    "Fake News [13], un corpus public collecté principalement sur Twitter entre "
    "janvier et juillet 2020. Le dataset comprend 10 700 échantillons répartis "
    "en un ensemble d'entraînement (6 420 exemples), de validation (2 140) et "
    "de test (2 140). Chaque échantillon est une courte publication accompagnée "
    "d'un label binaire : real (information vérifiée) ou fake (désinformation "
    "confirmée). L'annotation a été réalisée par des annotateurs formés "
    "s'appuyant sur des sources de référence comme l'OMS, le CDC, PubMed, "
    "Snopes et PolitiFact."
)

para(
    "La distribution des classes est remarquablement équilibrée : l'ensemble "
    "d'entraînement contient environ 52,3 % de news réelles et 47,7 % de fake "
    "news. Cet équilibre est intentionnel et simplifie considérablement la "
    "modélisation, en évitant les stratégies d'augmentation agressives qui "
    "introduiraient leurs propres biais statistiques."
)

add_figure(
    os.path.join(FIGURES, '01_class_distribution.png'),
    "Figure 1 : Distribution des classes dans l'ensemble d'entraînement."
)

para(
    "L'analyse de la longueur des textes (Figure 2) montre que la majorité "
    "des publications sont courtes — moins de 50 mots pour la plupart, ce qui "
    "est cohérent avec la limite de caractères de Twitter. Les distributions de "
    "longueur pour les deux classes se superposent largement, ce qui indique que "
    "la longueur seule n'est pas un signal discriminant suffisant et justifie "
    "le recours à des représentations lexicales et sémantiques plus riches."
)

add_figure(
    os.path.join(FIGURES, '02_text_length_distribution.png'),
    "Figure 2 : Distribution de la longueur des textes (nombre de mots) par classe."
)

heading("3.2 Prétraitement des données", level=2)

para(
    "Les textes bruts issus des réseaux sociaux nécessitent plusieurs étapes de "
    "normalisation avant de pouvoir être utilisés comme entrées pour les modèles. "
    "Notre pipeline de prétraitement applique les transformations suivantes dans "
    "l'ordre :"
)

steps = doc.add_paragraph()
steps.paragraph_format.left_indent = Cm(1)
steps.paragraph_format.space_after = Pt(4)
items = [
    "Suppression des URLs (http://, https://, www.) qui ne contiennent aucune valeur sémantique après tokenisation.",
    "Suppression des balises HTML résiduelles.",
    "Normalisation des mentions et hashtags : les @mentions sont supprimées, les #hashtags sont conservés sans le symbole '#'.",
    "Filtrage des caractères spéciaux : les caractères non alphanumériques sont remplacés par des espaces.",
    "Normalisation des espaces et mise en minuscules pour les modèles classiques.",
]
for i, item in enumerate(items, 1):
    p = doc.add_paragraph(f"{i}. {item}", style='List Number')
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph()
para(
    "Pour les modèles classiques, le texte nettoyé est ensuite transformé en "
    "une matrice de features TF-IDF, avec un vocabulaire de 15 000 termes, "
    "en utilisant des unigrammes et des bigrammes (ngram_range=(1,2)), une "
    "fréquence de document minimale de 2, une fréquence maximale de 95 %, et "
    "un scaling TF sublinéaire pour réduire la dominance des termes très "
    "fréquents. La matrice résultante a des dimensions de 6 420 × 15 000 pour "
    "l'ensemble d'entraînement."
)

para(
    "Nous extrayons également sept features linguistiques « artisanales » qui "
    "capturent des signaux de surface caractéristiques de la désinformation : "
    "nombre de mots, nombre de caractères, longueur moyenne des mots, nombre de "
    "points d'exclamation, de points d'interrogation, ratio de majuscules, et "
    "nombre de phrases. Ces features, insuffisantes prises isolément, contribuent "
    "à enrichir les représentations et servent d'aides à l'interprétabilité."
)

heading("3.3 Architectures des modèles", level=2)

para(
    "Nous entraînons quatre classifieurs classiques, tous opérant sur les "
    "représentations TF-IDF :"
)

classif_data = [
    ("SVM (LinearSVC)", "Séparateur à Vaste Marge, régularisation L2, C=1.0, poids de classe équilibrés."),
    ("Régression Logistique", "Régularisation L2, C=1.0, solveur lbfgs, 1000 itérations max."),
    ("Random Forest", "200 arbres, critère Gini, bootstrap sampling, profondeur max = 20."),
    ("XGBoost", "200 estimateurs, profondeur max = 6, learning rate = 0.1."),
]
for name, desc in classif_data:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"• {name} : ")
    r1.font.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.size = Pt(11)

doc.add_paragraph()
para(
    "Pour le modèle de deep learning, nous construisons une tête de classification "
    "sur le backbone pré-entraîné distilbert-base-uncased [5]. DistilBERT est "
    "une version compressée de BERT obtenue par distillation de connaissance : "
    "il conserve 97 % des capacités de compréhension linguistique de BERT tout "
    "en étant 40 % plus petit et 60 % plus rapide. La tête de classification "
    "consiste en deux couches fully-connected intercalées de Dropout (p=0.3), "
    "prenant en entrée la représentation du token [CLS] (768 dimensions) et "
    "produisant deux logits correspondant aux classes fake et real."
)

heading("3.4 Entraînement et fine-tuning", level=2)

para(
    "Les modèles classiques sont entraînés en une seule passe avec la bibliothèque "
    "scikit-learn, de manière déterministe et peu coûteuse en calcul (moins de "
    "30 secondes sur CPU standard). Une validation croisée à 5 plis est effectuée "
    "pour le SVM afin de confirmer la stabilité des métriques rapportées."
)

para(
    "Le fine-tuning de DistilBERT utilise l'optimiseur AdamW avec un taux de "
    "décroissance de poids de 0.01 et un learning rate initial de 2×10⁻⁵. Nous "
    "utilisons un scheduler linéaire avec une phase de warm-up couvrant 10 % des "
    "étapes totales d'entraînement. Les séquences d'entrée sont tronquées ou "
    "paddées à une longueur maximale de 128 tokens, ce qui couvre la grande "
    "majorité des tweets du dataset. L'entraînement est conduit sur 3 époques "
    "avec une taille de batch de 32. Les gradients sont clippés à 1.0 pour "
    "stabiliser l'entraînement."
)

para(
    "La Figure 3 montre les courbes d'apprentissage de DistilBERT. La loss "
    "d'entraînement diminue régulièrement de 0,36 à l'époque 1 jusqu'à 0,20 à "
    "l'époque 3, et la loss de validation suit une trajectoire parallèle sans "
    "diverger, ce qui indique une bonne généralisation. La précision de validation "
    "progresse de 91,9 % à 92,9 % sur les trois époques."
)

add_figure(
    os.path.join(FIGURES, '10_bert_training_curves.png'),
    "Figure 3 : Courbes d'apprentissage de DistilBERT (loss et précision sur 3 époques)."
)

heading("3.5 Évaluation et métriques", level=2)

para(
    "Tous les modèles sont évalués sur l'ensemble de validation à l'aide des "
    "métriques suivantes :"
)

metrics_rows = [
    ("Accuracy", "Proportion d'échantillons correctement classifiés."),
    ("F1-Score (macro)", "Moyenne harmonique de la précision et du rappel, pondérée par classe."),
    ("MCC", "Coefficient de Corrélation de Matthews, robuste aux classes déséquilibrées."),
    ("AUC-ROC", "Aire sous la courbe ROC, mesure la capacité discriminante globale."),
    ("Précision / Rappel", "Pour une vue détaillée des erreurs par type (FP vs FN)."),
]
for name, desc in metrics_rows:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"• {name} : ")
    r1.font.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.size = Pt(11)

doc.add_paragraph()
para(
    "Des matrices de confusion sont produites pour tous les modèles afin "
    "d'avoir une vue granulaire des distributions de faux positifs et faux négatifs."
)

heading("3.6 Méthodes d'explicabilité (XAI)", level=2)

para(
    "La transparence est une priorité de conception dans notre système. Nous "
    "employons deux techniques d'explication complémentaires."
)

para(
    "LIME (Local Interpretable Model-agnostic Explanations) [6] est une technique "
    "post-hoc et agnostique au modèle. Pour un texte donné, LIME perturbe l'entrée "
    "en masquant aléatoirement des mots, observe les changements de probabilités "
    "prédites, et ajuste un modèle linéaire sparse qui attribue des poids "
    "d'importance aux mots individuels. Les barres rouges indiquent les mots "
    "qui poussent la prédiction vers « fake » ; les barres vertes vers « real »."
)

para(
    "La visualisation des poids d'attention est appliquée spécifiquement à "
    "DistilBERT. Nous extrayons les poids d'attention moyennés de la dernière "
    "couche du transformeur, en nous concentrant sur le token [CLS] qui sert "
    "d'entrée au classificateur. Les mots recevant des poids d'attention élevés "
    "sont ceux auxquels le modèle s'est le plus attaché pour prendre sa décision."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 4. RÉSULTATS ET DISCUSSION
# ═══════════════════════════════════════════════════════════════════════════════
heading("4. Résultats et Discussion", level=1)

heading("4.1 Performance des modèles", level=2)

para(
    "Le Tableau 1 synthétise les performances de tous les modèles sur l'ensemble "
    "de validation, telles que visualisées dans la Figure 4."
)

add_table(
    ["Modèle", "Accuracy", "F1-Score", "Précision", "Rappel", "MCC", "AUC-ROC"],
    [
        ["SVM (LinearSVC) ★", "93,9 %", "0,939", "0,939", "0,939", "0,877", "—"],
        ["Régression Logistique", "92,6 %", "0,926", "0,926", "0,926", "0,852", "0,976"],
        ["Random Forest", "90,1 %", "0,901", "0,902", "0,901", "0,803", "0,969"],
        ["XGBoost", "90,4 %", "0,904", "0,906", "0,904", "0,810", "0,968"],
        ["DistilBERT", "92,9 %", "0,928", "0,929", "0,929", "0,857", "0,976"],
    ],
)

para(
    "Tableau 1 : Comparaison des performances de tous les modèles sur l'ensemble "
    "de validation (2 140 échantillons). ★ meilleur modèle.",
    italic=True, size=9
)

add_figure(
    os.path.join(FIGURES, '07_baseline_comparison.png'),
    "Figure 4 : Comparaison des modèles classiques sur Accuracy, F1, Précision et Rappel.",
    width=5.0
)

para(
    "Le SVM atteint la meilleure accuracy et le meilleur score F1 parmi tous les "
    "modèles (93,9 %), un résultat qui peut sembler contre-intuitif étant donné "
    "la bien plus grande complexité architecturale de DistilBERT. Cela confirme "
    "des observations récurrentes dans la littérature NLP : les modèles linéaires "
    "entraînés sur des représentations creuses de haute dimension restent très "
    "compétitifs sur la classification de textes courts, où le signal discriminant "
    "est essentiellement lexical [16]. La supériorité du SVM sur la Régression "
    "Logistique (1,3 points de pourcentage) reflète probablement son biais inductif "
    "à grande marge, particulièrement efficace quand le nombre de features (15 000 "
    "dimensions TF-IDF) dépasse celui des échantillons d'entraînement."
)

para(
    "La Régression Logistique, malgré sa simplicité, obtient le meilleur AUC-ROC "
    "parmi les modèles classiques (0,976), confirmant sa bonne calibration de "
    "probabilité. Les Random Forest et XGBoost, deux ensembles d'arbres de "
    "décision, accusent un retard d'environ 3,5 à 3,8 points de pourcentage "
    "sur les modèles linéaires. Cet écart est cohérent avec la tendance connue "
    "des ensembles d'arbres à rencontrer des difficultés dans les espaces de "
    "features creux et de haute dimension, où le sous-échantillonnage aléatoire "
    "des features à chaque nœud de division peut fréquemment sélectionner des "
    "features non informatives."
)

add_figure(
    os.path.join(FIGURES, 'roc_curves.png'),
    "Figure 5 : Courbes ROC pour les modèles baseline.",
    width=4.5
)

para(
    "DistilBERT atteint une précision de 92,9 % et une AUC-ROC de 0,976 — "
    "égalant la Régression Logistique en capacité discriminante tout en restant "
    "légèrement en dessous du SVM en précision brute. Le MCC de 0,857 le place "
    "en deuxième position après le SVM, confirmant une bonne généralisation. "
    "La matrice de confusion montre 1 048 vrais positifs, 939 vrais négatifs, "
    "72 faux positifs et 81 faux négatifs. L'asymétrie entre faux positifs et "
    "faux négatifs mérite attention : dans un contexte de santé publique, "
    "manquer de la désinformation véritable (faux négatifs) est sans doute "
    "plus coûteux que signaler à tort un contenu véridique (faux positifs)."
)

add_figure(
    os.path.join(FIGURES, '06_confusion_matrices_baselines.png'),
    "Figure 6 : Matrices de confusion pour SVM, Régression Logistique, Random Forest et XGBoost.",
    width=5.5
)

heading("4.2 Analyse de l'importance des features", level=2)

para(
    "L'analyse des coefficients de la Régression Logistique (Figure 7) fournit "
    "un éclairage direct sur les features lexicales qui distinguent les vraies "
    "nouvelles des fausses. Les meilleurs prédicteurs de news réelles incluent : "
    "« covid19 » (utilisé dans les hashtags officiels structurés), « rt » "
    "(retweet de comptes officiels), « cases », « testing », « data », "
    "« reported », « restrictions ». Ce vocabulaire reflète le langage de la "
    "surveillance épidémiologique et de la communication institutionnelle — "
    "factuel, numérique, géographiquement spécifique."
)

para(
    "À l'inverse, les meilleurs prédicteurs de fake news incluent : "
    "« coronavirus » (utilisé hors contexte institutionnel), « trump », "
    "« cure », « claim », « video », « china », « died ». Ce profil "
    "vocabulaire est caractéristique du contenu sensationnaliste et politisé "
    "qui attribue la pandémie à des acteurs spécifiques ou promeut des "
    "remèdes non vérifiés. La présence de « trump » comme fort indicateur "
    "de fake news reflète la politisation marquée du discours COVID-19 sur "
    "les réseaux sociaux en langue anglaise."
)

add_figure(
    os.path.join(FIGURES, '09_feature_importance_lr.png'),
    "Figure 7 : Top 20 coefficients TF-IDF de la Régression Logistique pour les classes REAL (gauche) et FAKE (droite).",
    width=5.5
)

heading("4.3 Analyse de l'explicabilité", level=2)

para(
    "Les explications LIME pour quatre exemples représentatifs (Figure 8) "
    "illustrent la logique décisionnelle du modèle au niveau de l'instance. "
    "Pour une publication correctement classifiée comme vraie (confiance : 99 %), "
    "LIME attribue les poids positifs les plus élevés à « data », « covid19 », "
    "« announcement » et « latest » — des mots signalant une communication "
    "officielle et factuelle. Pour une fake news correctement identifiée "
    "(confiance : 99 %), les features LIME dominantes sont des termes liés "
    "à des narratives de complot spécifiques, combinant des identificateurs "
    "ethniques ou nationaux avec « coronavirus »."
)

add_figure(
    os.path.join(FIGURES, 'lime_fake_1.png'),
    "Figure 8 : Explications LIME pour une prédiction de fake news.",
    width=4.5
)

para(
    "Ces résultats LIME soulignent une limitation importante des modèles "
    "purement data-driven : le classificateur a appris de fortes associations "
    "entre certains identificateurs ethniques ou nationaux et la classe fake, "
    "non parce que ces associations sont intrinsèquement correctes, mais parce "
    "que le dataset contient de nombreuses publications fabriquées exploitant "
    "des narratives de bouc-émissaire culturel. Ce résultat illustre le risque "
    "de propagation de biais et renforce la nécessité d'une revue humaine avant "
    "tout déploiement à grande échelle."
)

para(
    "La visualisation des poids d'attention (Figure 9) apporte une perspective "
    "complémentaire. Pour une vraie nouvelle, le token [CLS] distribue l'attention "
    "de manière diffuse sur plusieurs mots de contenu informatifs. Pour une fake "
    "news, l'attention se concentre de façon plus marquée sur les termes "
    "culturellement et contextuellement chargés — ce qui corrobore les "
    "observations faites avec LIME."
)

add_figure(
    os.path.join(FIGURES, '17_attention_visualization.png'),
    "Figure 9 : Visualisation des poids d'attention DistilBERT pour des exemples réels et faux.",
    width=5.5
)

heading("4.4 Analyse des erreurs", level=2)

para(
    "La distribution des niveaux de confiance révèle un pattern fortement "
    "bimodal : la grande majorité des prédictions, qu'elles soient correctes "
    "ou non, sont faites avec une confiance très élevée (P(fake) proche de 0 "
    "ou de 1). Les prédictions incorrectes sont plus uniformément réparties "
    "à des valeurs de confiance intermédiaires, ce qui suggère que l'incertitude "
    "du modèle est dans une certaine mesure calibrée."
)

para(
    "Le taux d'erreur en fonction de la longueur du texte montre une décroissance "
    "monotone claire : les textes de 3 à 13 mots ont un taux d'erreur d'environ "
    "11 %, tandis que les textes de 29 à 38 mots donnent un taux d'erreur "
    "inférieur à 6 %. Les textes très courts fournissent un contexte linguistique "
    "insuffisant pour que le modèle effectue des jugements corrects et confiants. "
    "Cette observation a une implication pratique directe : un système de "
    "déploiement devrait envisager d'appliquer un filtre de longueur qui oriente "
    "les textes très courts vers une revue humaine plutôt que de s'appuyer "
    "uniquement sur le classificateur automatique."
)

add_figure(
    os.path.join(FIGURES, 'final_results_heatmap.png'),
    "Figure 10 : Heatmap de performance de tous les modèles sur l'ensemble des métriques.",
    width=5.0
)

# ═══════════════════════════════════════════════════════════════════════════════
# 5. CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════════
heading("5. Conclusion et Perspectives", level=1)

para(
    "Ce travail a présenté une étude comparative complète de la détection "
    "automatique de fake news COVID-19, en combinant des baselines classiques "
    "avec un transformeur DistilBERT fine-tuné et une analyse d'explicabilité "
    "approfondie. Quatre conclusions principales se dégagent."
)

para(
    "Premièrement, le LinearSVC entraîné sur des features TF-IDF atteint la "
    "meilleure accuracy (93,9 %) et le meilleur MCC (0,877) de tous les modèles. "
    "Ce résultat démontre que les approches classiques bien calibrées restent "
    "très compétitives sur la détection de désinformation en textes courts. "
    "Il ne faut donc pas les écarter au profit des architectures Transformer "
    "sans justification empirique préalable."
)

para(
    "Deuxièmement, DistilBERT égale le meilleur modèle classique en AUC-ROC "
    "(0,976) tout en étant légèrement inférieur en précision brute. Ses "
    "représentations contextuelles offrent une meilleure calibration de "
    "probabilité et une interprétabilité supérieure via les mécanismes "
    "d'attention natifs."
)

para(
    "Troisièmement, LIME et la visualisation de l'attention révèlent que les "
    "modèles exploitent des indices linguistiques partiellement valables mais "
    "potentiellement biaisés — notamment la surreprésentation de termes "
    "ethniques et politiques comme prédicteurs de fake news. Cela constitue "
    "une préoccupation de fairness significative pour tout déploiement réel."
)

para(
    "Quatrièmement, l'analyse des erreurs identifie les textes très courts "
    "comme le principal mode d'échec, recommandant une architecture hybride "
    "avec intervention humaine pour les publications de moins de 15 mots."
)

para(
    "Plusieurs pistes méritent d'être explorées dans la continuité de ce "
    "travail. L'extension multimodale est la plus immédiate : les fake news "
    "sur les réseaux sociaux sont de plus en plus diffusées avec des images "
    "ou des vidéos, et l'intégration de features visuelles améliorerait "
    "vraisemblablement les performances. La généralisation multilingue "
    "constitue un autre axe important, étant donné que la désinformation "
    "COVID-19 était tout aussi présente en français, en arabe et en mandarin. "
    "Enfin, les implications de nos résultats d'explicabilité sur la mitigation "
    "des biais méritent une investigation dédiée, par exemple via de "
    "l'augmentation de données contrefactuelles."
)

# ═══════════════════════════════════════════════════════════════════════════════
# RÉFÉRENCES
# ═══════════════════════════════════════════════════════════════════════════════
heading("Références", level=1)

refs = [
    "[1] Organisation Mondiale de la Santé, « Novel Coronavirus (2019-nCoV) Situation Report – 13 », OMS, Genève, 2020.",
    "[2] K. Sharma, F. Qian, H. Jiang, N. Ruchansky, M. Zhang et Y. Liu, « Combating Fake News: A Survey on Identification and Mitigation Techniques », ACM TIST, vol. 10, no 3, pp. 1–42, 2019.",
    "[3] V. Pérez-Rosas, B. Kleinberg, A. Lefevre et R. Mihalcea, « Automatic Detection of Fake News », Proc. COLING 2018, pp. 3391–3401.",
    "[4] J. Devlin, M.-W. Chang, K. Lee et K. Toutanova, « BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding », Proc. NAACL-HLT 2019, pp. 4171–4186.",
    "[5] V. Sanh, L. Debut, J. Chaumond et T. Wolf, « DistilBERT, a Distilled Version of BERT: Smaller, Faster, Cheaper and Lighter », arXiv:1910.01108, 2019.",
    "[6] M. T. Ribeiro, S. Singh et C. Guestrin, « 'Why Should I Trust You?': Explaining the Predictions of Any Classifier », Proc. KDD 2016, pp. 1135–1144.",
    "[7] S. M. Lundberg et S.-I. Lee, « A Unified Approach to Interpreting Model Predictions », NeurIPS 2017, pp. 4765–4774.",
    "[8] D. Pomerleau et D. Rao, « Fake News Challenge Stage 1 (FNC-1): Stance Detection », FakeNewsChallenge.org, 2017.",
    "[9] N. Rashkin et al., « Truth of Varying Shades: Analyzing Language in Fake News and Political Fact-Checking », Proc. EMNLP 2017, pp. 2931–2937.",
    "[10] S. Volkova et al., « Separating Facts from Fiction: Linguistic Models to Classify Suspicious and Trusted News Posts on Twitter », Proc. ACL 2017, pp. 647–652.",
    "[11] S. Kula et al., « Application of the BERT-Based Architecture in Fake News Detection », Proc. ICCS 2020, LNCS vol. 12140, pp. 239–248.",
    "[12] M. H. Goldani, S. Momtazi et R. Safabakhsh, « Detecting Fake News with Capsule Neural Networks », Applied Soft Computing, vol. 101, 2021.",
    "[13] P. Patwa et al., « Fighting an Infodemic: COVID-19 Fake News Dataset », Proc. Constraint Workshop @ AAAI 2021, pp. 21–29.",
    "[14] D. Popat et al., « Where the Truth Lies: Explaining the Credibility of Emerging Claims on the Web », Proc. WWW 2018, pp. 1047–1056.",
    "[15] S. Jain et B. C. Wallace, « Attention is not Explanation », Proc. NAACL-HLT 2019, pp. 3543–3556.",
    "[16] T. Joachims, « Text Categorization with Support Vector Machines: Learning with Many Relevant Features », Proc. ECML 1998, pp. 137–142.",
    "[17] P. Nakov et al., « Overview of the CLEF-2021 CheckThat! Lab », Proc. CLEF 2021, pp. 264–291.",
]

for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(10)

# ── Save ─────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Saved : {OUT}")
pages_estimate = len(doc.paragraphs) // 8
print(f"Estimated paragraphs : {len(doc.paragraphs)}")
print(f"Figures embedded : {sum(1 for p in doc.paragraphs if 'Figure' in p.text and p.text.startswith('Figure'))}")
